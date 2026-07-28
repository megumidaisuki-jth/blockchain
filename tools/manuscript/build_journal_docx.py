from __future__ import annotations

import argparse
import posixpath
import re
import shutil
import sys
import zipfile
from pathlib import Path

VENDOR_DIR = Path(__file__).resolve().parent / "_vendor"
if VENDOR_DIR.exists():
    sys.path.insert(0, str(VENDOR_DIR))

from latex2mathml.converter import convert as latex_to_mathml
from lxml import etree
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


STYLE_FALLBACKS = {
    "人名": "Normal",
    "地名": "Normal",
    "摘要": "Normal",
    "公式": "Normal",
    "图注": "Normal",
    "表题": "Normal",
    "文献": "Normal",
    "文献文": "Normal",
    "作者简介：": "Normal",
}


def style_name(doc: Document, preferred: str) -> str:
    names = {s.name for s in doc.styles}
    return preferred if preferred in names else STYLE_FALLBACKS.get(preferred, "Normal")


def set_cols(section, count: int, space_twips: int = 426) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    if count == 1:
        cols.attrib.pop(qn("w:num"), None)
        cols.set(qn("w:space"), "720")
    else:
        cols.set(qn("w:num"), str(count))
        cols.set(qn("w:space"), str(space_twips))


def set_cell_width(cell, twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")


def set_fixed_table_geometry(table, widths: list[int]) -> None:
    """Set table, grid, and cell widths consistently in DXA."""
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def set_cell_margins(cell, *, top: int = 0, right: int = 0, bottom: int = 0, left: int = 0) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("right", right), ("bottom", bottom), ("left", left)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, attrs in edges.items():
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in attrs.items():
            element.set(qn(f"w:{key}"), str(value))


def three_line_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(
                cell,
                top={"val": "nil"},
                bottom={"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"},
                insideH={"val": "nil"},
                insideV={"val": "nil"},
            )
            if row_idx == 0:
                set_cell_border(cell, top={"val": "single", "sz": "8", "color": "000000"}, bottom={"val": "single", "sz": "4", "color": "000000"})
            if row_idx == len(table.rows) - 1:
                set_cell_border(cell, bottom={"val": "single", "sz": "8", "color": "000000"})


GREEK = {
    r"\tau": "τ", r"\alpha": "α", r"\beta": "β", r"\Gamma": "Γ",
    r"\Sigma": "Σ", r"\Delta": "Δ", r"\lambda": "λ", r"\rho": "ρ",
    r"\pi": "π", r"\xi": "ξ", r"\eta": "η", r"\theta": "θ",
    r"\varepsilon": "ε", r"\delta": "δ", r"\ell": "ℓ",
}


MML2OMML_CANDIDATES = (
    Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"),
    Path(r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL"),
)


def load_mml2omml_transform():
    for path in MML2OMML_CANDIDATES:
        if path.exists():
            return etree.XSLT(etree.parse(str(path)))
    raise FileNotFoundError(
        "MML2OMML.XSL not found. Install Microsoft Word or provide its Office16 transform."
    )


MML2OMML = load_mml2omml_transform()


def latex_omml(formula: str):
    """Convert one display formula from LaTeX to native, editable Word OMML."""
    latex = re.sub(r"\s*\\tag\{[^}]+\}\s*", "", formula).strip()
    # latex2mathml 3.81.0 does not recognize \qquad and otherwise emits the
    # literal control word.  Two \quad tokens preserve the intended spacing.
    latex = latex.replace(r"\qquad", r"\quad\quad")
    mathml = latex_to_mathml(latex)
    transformed = MML2OMML(etree.fromstring(mathml.encode("utf-8")))
    root = transformed.getroot()
    if root is None:
        raise ValueError(f"empty OMML conversion for: {latex}")
    return root


def unicode_math(text: str) -> str:
    text = text.strip()
    text = re.sub(r"\\tag\{[^}]+\}", "", text)
    replacements = {
        r"\mathcal H": "𝓗", r"\mathcal I": "𝓘", r"\mathcal S": "𝒮", r"\mathcal A": "𝒜",
        r"\mathbb Z": "ℤ", r"\mathbb R": "ℝ", r"\boldsymbol 1": "𝟙",
        r"\operatorname{Cov}": "Cov", r"\operatorname{diag}": "diag",
        r"\mathrm E": "E", r"\Pr": "Pr", r"\Rightarrow": "⇒", r"\longrightarrow": "→",
        r"\to": "→", r"\ge": "≥", r"\le": "≤", r"\ne": "≠", r"\in": "∈",
        r"\prod": "∏", r"\sum": "∑", r"\inf": "inf", r"\min": "min", r"\max": "max",
        r"\cdot": "·", r"\times": "×", r"\pm": "±", r"\infty": "∞", r"\ldots": "…",
        r"\{": "{", r"\}": "}", r"\,": " ", r"\;": " ", r"\quad": " ", r"\\": " ",
    }
    for src, dst in {**GREEK, **replacements}.items():
        text = text.replace(src, dst)
    text = re.sub(r"\\(?:mathrm|boldsymbol|operatorname)\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\(?:mathcal|mathbb)\{([^{}]+)\}", r"\1", text)
    text = text.replace("{", "(").replace("}", ")")
    text = re.sub(r"\s+", " ", text).strip()
    return text


INLINE_MATH = re.compile(r"\\\((.+?)\\\)")
STRONG = re.compile(r"\*\*(.+?)\*\*")


def add_inline_runs(paragraph, text: str) -> None:
    pos = 0
    tokens = []
    for match in re.finditer(r"\\\((.+?)\\\)|\*\*(.+?)\*\*", text):
        if match.start() > pos:
            tokens.append(("text", text[pos:match.start()]))
        if match.group(1) is not None:
            # Keep inline mathematics as native, editable Word OMML.  The
            # previous Unicode fallback exposed source-like underscores and
            # braces in Word and was vulnerable to font-dependent garbling.
            tokens.append(("math", match.group(1)))
        else:
            tokens.append(("bold", match.group(2)))
        pos = match.end()
    if pos < len(text):
        tokens.append(("text", text[pos:]))
    for kind, value in tokens:
        if kind == "math":
            paragraph._p.append(latex_omml(value))
            continue
        run = paragraph.add_run(value)
        if kind == "bold":
            run.bold = True


def add_formula(doc: Document, formula: str) -> None:
    tag_match = re.search(r"\\tag\{([^}]+)\}", formula)
    tag = f"({tag_match.group(1)})" if tag_match else ""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = table.rows[0].cells
    # The journal body is two-column.  Keep the equation and its number inside
    # one 4,602-DXA column instead of inheriting a full-page 9,000-DXA grid.
    set_fixed_table_geometry(table, [4080, 500])
    for cell in (left, right):
        set_cell_margins(cell)
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})
    p = left.paragraphs[0]
    p.style = style_name(doc, "公式")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p._p.append(latex_omml(formula))
    p2 = right.paragraphs[0]
    p2.style = style_name(doc, "公式")
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.add_run(tag)


def add_full_width_section(doc: Document) -> None:
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    set_cols(sec, 1)


def return_to_two_columns(doc: Document) -> None:
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    set_cols(sec, 2)


def clear_reference_body(doc: Document) -> None:
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def remove_orphan_media(docx_path: Path) -> tuple[int, int]:
    """Drop template media files that are no longer referenced by any OOXML part."""
    relationship_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    with zipfile.ZipFile(docx_path, "r") as source:
        names = source.namelist()
        referenced: set[str] = set()
        rewritten_relationships: dict[str, bytes] = {}
        for rel_name in (name for name in names if name.endswith(".rels")):
            root = etree.fromstring(source.read(rel_name))
            source_dir = posixpath.dirname(posixpath.dirname(rel_name))
            source_part = posixpath.join(
                source_dir, posixpath.basename(rel_name)[: -len(".rels")]
            )
            used_ids: set[str] = set()
            if source_part in names and source_part.endswith(".xml"):
                source_root = etree.fromstring(source.read(source_part))
                used_ids = {
                    value
                    for element in source_root.iter()
                    for value in element.attrib.values()
                    if value.startswith("rId")
                }
            for relation in list(root.findall(f"{{{relationship_ns}}}Relationship")):
                if relation.get("TargetMode") == "External":
                    continue
                target = relation.get("Target", "")
                resolved = posixpath.normpath(posixpath.join(source_dir, target))
                if resolved.startswith("word/media/"):
                    if relation.get("Id") in used_ids:
                        referenced.add(resolved)
                    else:
                        root.remove(relation)
            rewritten_relationships[rel_name] = etree.tostring(
                root, xml_declaration=True, encoding="UTF-8", standalone=True
            )
        media_names = {name for name in names if name.startswith("word/media/")}
        orphaned = media_names - referenced
        temp_path = docx_path.with_suffix(docx_path.suffix + ".tmp")
        with zipfile.ZipFile(temp_path, "w", compression=zipfile.ZIP_DEFLATED) as target_zip:
            for item in source.infolist():
                if item.filename not in orphaned:
                    target_zip.writestr(
                        item,
                        rewritten_relationships.get(item.filename, source.read(item.filename)),
                    )
    temp_path.replace(docx_path)
    return len(referenced), len(orphaned)


def parse_markdown(md_path: Path, reference: Path, output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(reference, output)
    doc = Document(output)
    clear_reference_body(doc)
    lines = md_path.read_text(encoding="utf-8").splitlines()
    title = next(
        (line[2:].strip() for line in lines if line.startswith("# ")),
        "原子路由下超图支付通道网络的首次耗尽时间",
    )
    keyword_line = next(
        (line for line in lines if line.startswith("**关键词：**")),
        "",
    )
    keywords = keyword_line.removeprefix("**关键词：**").strip().replace("；", "; ")
    first = doc.sections[0]
    set_cols(first, 1)
    first.different_first_page_header_footer = True
    first.header.is_linked_to_previous = False
    first.first_page_header.paragraphs[0].text = ""
    hp = first.header.paragraphs[0]
    hp.text = f"\t[作者]等：{title}\t"
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    i = 0
    front = True
    figure_count = 0
    table_count = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line or line.startswith("<!--"):
            i += 1
            continue
        if line.startswith("!["):
            m = re.match(r"!\[(.+?)\]\((.+?)\)", line)
            if m:
                caption, image_path = m.groups()
                add_full_width_section(doc)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(Path(image_path)), width=Cm(16.5))
                cp = doc.add_paragraph(style=style_name(doc, "图注"))
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.add_run(caption)
                figure_count += 1
                return_to_two_columns(doc)
            i += 1
            continue
        if line == r"\[":
            formula_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != r"\]":
                formula_lines.append(lines[i].strip())
                i += 1
            add_formula(doc, " ".join(formula_lines))
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[i + 1].strip()):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                parts = [p.strip() for p in lines[i].strip().strip("|").split("|")]
                rows.append(parts)
                i += 1
            rows = [rows[0]] + rows[2:]
            add_full_width_section(doc)
            table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
            table.autofit = True
            for ri, row in enumerate(rows):
                for ci, value in enumerate(row):
                    p = table.cell(ri, ci).paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_inline_runs(p, value)
                    if ri == 0:
                        for run in p.runs:
                            run.bold = True
            three_line_table(table)
            table_count += 1
            return_to_two_columns(doc)
            continue
        if line.startswith("# "):
            p = doc.add_paragraph(style="Heading 1")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line[2:].strip())
            i += 1
            continue
        if line.startswith("## "):
            heading = line[3:].strip()
            if heading == "0　引言" and front:
                sec = doc.add_section(WD_SECTION.CONTINUOUS)
                set_cols(sec, 2)
                front = False
            p = doc.add_paragraph(style="Heading 2")
            add_inline_runs(p, heading)
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline_runs(p, line[4:].strip())
            i += 1
            continue
        if re.match(r"^\[\d+\]\s", line):
            p = doc.add_paragraph(style=style_name(doc, "文献文"))
            add_inline_runs(p, line)
            i += 1
            continue
        if line.startswith("（1. ") or line.startswith("(1. "):
            p = doc.add_paragraph(style=style_name(doc, "地名"))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, line)
            i += 1
            continue
        if line.startswith("**[作者") or line.startswith("**[AUTHOR"):
            p = doc.add_paragraph(style=style_name(doc, "人名"))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, line)
            i += 1
            continue
        if line.startswith("**摘  要：") or line.startswith("**关键词：") or line.startswith("**中图分类号：") or line.startswith("**Abstract:") or line.startswith("**Keywords:"):
            p = doc.add_paragraph(style=style_name(doc, "摘要"))
            add_inline_runs(p, line)
            i += 1
            continue
        if line.startswith("1. ") or line.startswith("2. ") or line.startswith("3. ") or line.startswith("4. "):
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.first_line_indent = Cm(-0.74)
            add_inline_runs(p, line)
            i += 1
            continue
        if line.startswith("**表"):
            p = doc.add_paragraph(style=style_name(doc, "表题"))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, line)
            i += 1
            continue
        if line.startswith("**定理") or line.startswith("**证明"):
            p = doc.add_paragraph(style="Normal")
            add_inline_runs(p, line)
            i += 1
            continue
        if line.startswith("代码仓库："):
            p = doc.add_paragraph(style="Normal")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            add_inline_runs(p, line)
            i += 1
            continue
        p_style = "Normal"
        if line == "参考文献":
            p_style = style_name(doc, "文献")
        p = doc.add_paragraph(style=p_style)
        add_inline_runs(p, line)
        i += 1

    for section in doc.sections:
        section.top_margin = Cm(2.41)
        section.bottom_margin = Cm(2.41)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.header_distance = Pt(49.6)
        section.footer_distance = Pt(38.25)
    core = doc.core_properties
    core.title = title
    core.subject = "《通信学报》格式论文稿"
    core.author = "[待作者填写]"
    core.keywords = keywords
    doc.save(output)
    referenced_media, removed_media = remove_orphan_media(output)
    print(f"built={output}")
    print(f"sections={len(doc.sections)} figures={figure_count} tables={table_count}")
    print(f"referenced_media={referenced_media} removed_orphan_media={removed_media}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--markdown", type=Path, required=True)
    parser.add_argument("--reference", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    parse_markdown(args.markdown.resolve(), args.reference.resolve(), args.output.resolve())


if __name__ == "__main__":
    main()
